"""
Standup/Workshop-protokoll -> Word (.docx)

Krav:
  pip install python-docx

K칬r:
  python standup_to_word.py

Skapar:
  ./protokoll_<Team>_<YYYY-MM-DD>.docx
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _safe_filename(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^A-Za-z0-9_\-칀츿칐친칛칬]", "", s)
    return s or "Team"


def _split_lines(text: str) -> list[str]:
    """Split text into non-empty stripped lines."""
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if line:
            lines.append(line)
    return lines


def _map_status_to_emoji(status_label: str) -> str:
    # status_label is one of the dropdown values
    mapping = {
        "P친 sp친r (游릭)": "游릭 P친 sp친r",
        "Lite efter (游리)": "游리 Lite efter",
        "Beh칬ver hj칛lp (游댮)": "游댮 Beh칬ver hj칛lp",
    }
    return mapping.get(status_label, "游릭 P친 sp친r")


def create_word_doc(
    team: str,
    datum: str,
    deltagare: str,
    vad_vi_jobbat_med: str,
    hinder: str,
    status_label: str,
    nasta_steg: str,
    output_dir: Path | None = None,
) -> Path:
    output_dir = output_dir or Path("Protokoll")
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_paragraph("Workshop-protokoll")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    # Header fields
    p = doc.add_paragraph()
    p.add_run("Team: ").bold = True
    p.add_run(team.strip() or "-")

    p = doc.add_paragraph()
    p.add_run("Datum: ").bold = True
    p.add_run(datum.strip() or "-")

    p = doc.add_paragraph()
    p.add_run("Deltagare: ").bold = True
    p.add_run(deltagare.strip() or "-")

    doc.add_paragraph("")

    # Sections
    def add_section(title_text: str, items_text: str):
        h = doc.add_paragraph(title_text)
        h.runs[0].bold = True
        for item in _split_lines(items_text):
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph("")

    add_section("Vad vi jobbade med:", vad_vi_jobbat_med)
    add_section("Hinder vi st칬tte p친:", hinder)

    # Status
    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    p.add_run(_map_status_to_emoji(status_label))
    doc.add_paragraph("")

    add_section("N칛sta steg:", nasta_steg)

    safe_team = _safe_filename(team)
    filename = f"protokoll_{safe_team}_{datum}.docx"
    outpath = output_dir / filename
    doc.save(outpath)
    return outpath


def main():
    root = tk.Tk()
    root.title("Protokoll -> Word")

    # Layout
    frm = ttk.Frame(root, padding=12)
    frm.grid(row=0, column=0, sticky="nsew")
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)
    frm.columnconfigure(1, weight=1)

    # Variables
    team_var = tk.StringVar()
    datum_var = tk.StringVar(value=str(date.today()))
    deltagare_var = tk.StringVar()
    status_var = tk.StringVar(value="P친 sp친r (游릭)")

    def add_label(row: int, text: str):
        ttk.Label(frm, text=text).grid(row=row, column=0, sticky="w", padx=(0, 10), pady=4)

    def add_entry(row: int, var: tk.StringVar):
        e = ttk.Entry(frm, textvariable=var)
        e.grid(row=row, column=1, sticky="ew", pady=4)
        return e

    add_label(0, "Team")
    add_entry(0, team_var)

    add_label(1, "Datum (YYYY-MM-DD)")
    add_entry(1, datum_var)

    add_label(2, "Deltagare (komma-separerat)")
    add_entry(2, deltagare_var)

    # Multi-line fields
    def add_text(row: int, title: str, height: int = 5) -> tk.Text:
        add_label(row, title)
        txt = tk.Text(frm, height=height, wrap="word")
        txt.grid(row=row, column=1, sticky="ew", pady=4)
        return txt

    vad_txt = add_text(3, "Vad vi jobbade med (en punkt per rad)", height=6)
    hinder_txt = add_text(4, "Hinder vi st칬tte p친 (en punkt per rad)", height=6)

    add_label(5, "Status")
    status_cb = ttk.Combobox(
        frm,
        textvariable=status_var,
        values=["P친 sp친r (游릭)", "Lite efter (游리)", "Beh칬ver hj칛lp (游댮)"],
        state="readonly",
    )
    status_cb.grid(row=5, column=1, sticky="w", pady=4)

    nasta_txt = add_text(6, "N칛sta steg (en punkt per rad)", height=6)

    def on_generate():
        team = team_var.get().strip()
        datum = datum_var.get().strip()
        deltagare = deltagare_var.get().strip()

        # Enkel validering p친 datumformat
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", datum):
            messagebox.showerror("Fel", "Datum m친ste vara i formatet YYYY-MM-DD (t.ex. 2026-01-20).")
            return

        try:
            outpath = create_word_doc(
                team=team or "Teamnamn",
                datum=datum,
                deltagare=deltagare or "-",
                vad_vi_jobbat_med=vad_txt.get("1.0", "end"),
                hinder=hinder_txt.get("1.0", "end"),
                status_label=status_var.get(),
                nasta_steg=nasta_txt.get("1.0", "end"),
            )
        except Exception as e:
            messagebox.showerror("Fel", f"Kunde inte skapa dokumentet:\n{e}")
            return

        messagebox.showinfo("Klart", f"Word-dokument skapat:\n{outpath}")

    btns = ttk.Frame(frm)
    btns.grid(row=7, column=0, columnspan=2, sticky="e", pady=(10, 0))
    ttk.Button(btns, text="Skapa Word-dokument", command=on_generate).grid(row=0, column=0)

    root.mainloop()


if __name__ == "__main__":
    main()

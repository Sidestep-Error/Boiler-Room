"""
Standup/Workshop-protokoll -> Word (.docx) + Settings + GitHub upload

- GUI-formulär (tkinter)
- Skapar .docx (python-docx)
- Settings-fil (settings.json): output_dir + senast använda Team + GitHub-konfiguration
- Vid uppstart: beräknar förväntat filnamn baserat på (last_team + dagens datum).
  Om filen finns: fråga om den ska öppnas (befintlig) istället.
- GitHub: klona repo om det saknas lokalt, fråga branch, checkout, kopiera fil, commit & push.

Krav:
  pip install python-docx
  Git installerat och fungerande auth mot GitHub (SSH rekommenderas).
"""

from __future__ import annotations

import json
import os
import re
import sys
import subprocess
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ----------------------------
# Settings
# ----------------------------

SETTINGS_FILENAME = "settings.json"


@dataclass
class GitHubSettings:
    repo_url: str = ""
    local_repos_base: str = "~/dev"
    repo_subdir: str = "docs/protokoll"
    default_branch: str = "main"
    commit_prefix: str = "Lägg till protokoll"


@dataclass
class AppSettings:
    output_dir: str = "./output"
    last_team: str = "Teamnamn"
    github: GitHubSettings = field(default_factory=GitHubSettings)


def load_settings(settings_path: Path) -> AppSettings:
    if not settings_path.exists():
        return AppSettings()

    try:
        data = json.loads(settings_path.read_text(encoding="utf-8"))
    except Exception:
        return AppSettings()

    gh = data.get("github", {}) if isinstance(data, dict) else {}
    github = GitHubSettings(
        repo_url=str(gh.get("repo_url", "")),
        local_repos_base=str(gh.get("local_repos_base", "~/dev")),
        repo_subdir=str(gh.get("repo_subdir", "docs/protokoll")),
        default_branch=str(gh.get("default_branch", "main")),
        commit_prefix=str(gh.get("commit_prefix", "Lägg till protokoll")),
    )

    return AppSettings(
        output_dir=str(data.get("output_dir", "./output")),
        last_team=str(data.get("last_team", "Teamnamn")),
        github=github,
    )


def save_settings(settings_path: Path, settings: AppSettings) -> None:
    payload = {
        "output_dir": settings.output_dir,
        "last_team": settings.last_team,
        "github": {
            "repo_url": settings.github.repo_url,
            "local_repos_base": settings.github.local_repos_base,
            "repo_subdir": settings.github.repo_subdir,
            "default_branch": settings.github.default_branch,
            "commit_prefix": settings.github.commit_prefix,
        },
    }
    settings_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


# ----------------------------
# Utilities
# ----------------------------

def safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^A-Za-z0-9_\-ÅÄÖåäö]", "", s)
    return s or "Team"


def planned_docx_path(team: str, datum: str, output_dir: Path) -> Path:
    output_dir = output_dir.expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"protokoll_{safe_filename(team)}_{datum}.docx"
    return output_dir / filename


def open_file_with_default_app(path: Path) -> None:
    path = path.expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(path)

    if sys.platform.startswith("win"):
        os.startfile(str(path))  # type: ignore[attr-defined]
    elif sys.platform == "darwin":
        subprocess.run(["open", str(path)], check=True)
    else:
        subprocess.run(["xdg-open", str(path)], check=True)


def split_lines(text: str) -> list[str]:
    return [ln.strip() for ln in (text or "").splitlines() if ln.strip()]


def map_status_to_emoji(status_label: str) -> str:
    mapping = {
        "På spår (🟢)": "🟢 På spår",
        "Lite efter (🟡)": "🟡 Lite efter",
        "Behöver hjälp (🔴)": "🔴 Behöver hjälp",
    }
    return mapping.get(status_label, "🟢 På spår")


# ----------------------------
# Word generation
# ----------------------------

def load_existing_docx(docx_path: Path) -> dict[str, str]:
    """Läs ett befintligt Word-dokument och extrahera innehållet."""
    doc = Document(docx_path)

    result = {
        "team": "",
        "datum": "",
        "deltagare": "",
        "vad_vi_jobbat_med": "",
        "hinder": "",
        "status": "På spår (🟢)",
        "nasta_steg": "",
    }

    # Samla all text från dokumentet
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()

        # Skippa tomma rader
        if not text:
            continue

        # Skippa titeln
        if text == "Standup / Workshop-protokoll":
            continue

        # Läs metadata
        if text.startswith("Team:"):
            result["team"] = text.replace("Team:", "").strip()
            current_section = None
        elif text.startswith("Datum:"):
            result["datum"] = text.replace("Datum:", "").strip()
            current_section = None
        elif text.startswith("Deltagare:"):
            result["deltagare"] = text.replace("Deltagare:", "").strip()
            current_section = None
        elif text.startswith("Status:"):
            status_text = text.replace("Status:", "").strip()
            # Mappa tillbaka emoji till combo-box värde
            if "🟢" in status_text:
                result["status"] = "På spår (🟢)"
            elif "🟡" in status_text:
                result["status"] = "Lite efter (🟡)"
            elif "🔴" in status_text:
                result["status"] = "Behöver hjälp (🔴)"
            current_section = None
        # Identifiera sektioner
        elif text == "Vad vi jobbade med:":
            current_section = "vad_vi_jobbat_med"
        elif text == "Hinder vi stötte på:":
            current_section = "hinder"
        elif text == "Nästa steg:":
            current_section = "nasta_steg"
        # Lägg till innehåll till den aktiva sektionen
        elif current_section:
            # Ta bort eventuella bullet-tecken som Word lägger till
            cleaned_text = text.lstrip("•\t -")
            if result[current_section]:
                result[current_section] += "\n"
            result[current_section] += cleaned_text

    return result


def create_word_doc_to_path(
    outpath: Path,
    team: str,
    datum: str,
    deltagare: str,
    vad_vi_jobbat_med: str,
    hinder: str,
    status_label: str,
    nasta_steg: str,
) -> Path:
    outpath = outpath.expanduser().resolve()
    outpath.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = doc.add_paragraph("Boiler Room-protokoll")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Team: ").bold = True
    p.add_run(team.strip() or "-")

    p = doc.add_paragraph()
    p.add_run("Datum: ").bold = True
    p.add_run(datum.strip() or "-")

    p = doc.add_paragraph()
    p.add_run("Deltagare: ").bold = True
    p.add_run(deltagare.strip() or "-")

    doc.add_paragraph("")

    def add_section(title_text: str, items_text: str):
        h = doc.add_paragraph(title_text)
        h.runs[0].bold = True
        for item in split_lines(items_text):
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph("")

    add_section("Vad vi jobbade med:", vad_vi_jobbat_med)
    add_section("Hinder vi stötte på:", hinder)

    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    p.add_run(map_status_to_emoji(status_label))
    doc.add_paragraph("")

    add_section("Nästa steg:", nasta_steg)

    doc.save(outpath)
    return outpath


# ----------------------------
# GitHub / git workflow
# ----------------------------

def run_git(args: list[str], cwd: Path | None = None) -> str:
    proc = subprocess.run(
        ["git", *args],
        cwd=str(cwd) if cwd else None,
        text=True,
        capture_output=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(
            f"Git-kommando misslyckades: git {' '.join(args)}\n"
            f"STDOUT:\n{proc.stdout}\n"
            f"STDERR:\n{proc.stderr}"
        )
    return proc.stdout.strip()


def ensure_repo_cloned(repo_url: str, base_dir: Path, repo_name: str | None = None) -> Path:
    base_dir = base_dir.expanduser().resolve()
    base_dir.mkdir(parents=True, exist_ok=True)

    if not repo_url.strip():
        raise RuntimeError("GitHub repo_url saknas i settings.json (github.repo_url).")

    if not repo_name:
        name = repo_url.rstrip("/").split("/")[-1]
        if name.endswith(".git"):
            name = name[:-4]
        repo_name = name

    repo_dir = (base_dir / repo_name).resolve()

    if (repo_dir / ".git").exists():
        return repo_dir

    if repo_dir.exists() and any(repo_dir.iterdir()):
        raise RuntimeError(
            f"Mappen finns redan men är inte ett git-repo: {repo_dir}\n"
            f"Rensa mappen eller välj annan github.local_repos_base."
        )

    run_git(["clone", repo_url, str(repo_dir)], cwd=base_dir)
    return repo_dir


def prompt_branch(root: tk.Tk, default: str = "main") -> str | None:
    branch = simpledialog.askstring(
        "Välj branch",
        f"Vilken branch vill du committa till?\n(t.ex. {default}, develop, feature/protokoll)\n",
        initialvalue=default,
        parent=root,
    )
    if branch is None:
        return None
    branch = branch.strip() or default

    if re.search(r"[\s~^:?*\[\]\\]", branch):
        messagebox.showerror("Ogiltigt branch-namn", "Branch-namnet innehåller ogiltiga tecken.")
        return None

    return branch


def checkout_branch(repo_dir: Path, branch: str) -> None:
    repo_dir = repo_dir.resolve()
    run_git(["fetch", "--all", "--prune"], cwd=repo_dir)

    local = run_git(["branch", "--list", branch], cwd=repo_dir)
    if local.strip():
        run_git(["checkout", branch], cwd=repo_dir)
        return

    remote = run_git(["branch", "-r", "--list", f"origin/{branch}"], cwd=repo_dir)
    if remote.strip():
        run_git(["checkout", "-t", f"origin/{branch}"], cwd=repo_dir)
        return

    origin_head = run_git(["symbolic-ref", "refs/remotes/origin/HEAD"], cwd=repo_dir)
    default_remote_branch = origin_head.split("/")[-1] if origin_head else "main"
    run_git(["checkout", "-b", branch, f"origin/{default_remote_branch}"], cwd=repo_dir)


def copy_docx_into_repo(docx_path: Path, repo_dir: Path, repo_subdir: str) -> Path:
    target_dir = (repo_dir / repo_subdir).resolve()
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / docx_path.name
    target_path.write_bytes(docx_path.read_bytes())
    return target_path


def commit_and_push(repo_dir: Path, file_path_in_repo: Path, commit_prefix: str) -> None:
    rel = file_path_in_repo.relative_to(repo_dir)
    run_git(["add", str(rel)], cwd=repo_dir)

    status = run_git(["status", "--porcelain"], cwd=repo_dir)
    if not status.strip():
        return

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    commit_msg = f"{commit_prefix}: {file_path_in_repo.stem} ({timestamp})"
    run_git(["commit", "-m", commit_msg], cwd=repo_dir)

    try:
        run_git(["push"], cwd=repo_dir)
    except RuntimeError:
        current_branch = run_git(["rev-parse", "--abbrev-ref", "HEAD"], cwd=repo_dir).strip()
        run_git(["push", "-u", "origin", current_branch], cwd=repo_dir)


def upload_docx_flow(root: tk.Tk, docx_path: Path, gh: GitHubSettings) -> tuple[Path, str]:
    repo_dir = ensure_repo_cloned(
        repo_url=gh.repo_url,
        base_dir=Path(gh.local_repos_base),
        repo_name=None,
    )

    branch = prompt_branch(root, default=gh.default_branch)
    if branch is None:
        raise RuntimeError("Avbrutet: ingen branch vald.")

    checkout_branch(repo_dir, branch)
    target_path = copy_docx_into_repo(docx_path, repo_dir, repo_subdir=gh.repo_subdir)
    commit_and_push(repo_dir, target_path, commit_prefix=gh.commit_prefix)

    return target_path, branch


# ----------------------------
# GUI app
# ----------------------------

def ask_file_exists_action(parent, filepath: Path) -> str | None:
    """
    Visa en custom dialog med tre alternativ: Ladda in, Uppdatera, Avbryt.
    Returnerar: 'load', 'update', eller None (för avbryt)
    """
    dialog = tk.Toplevel(parent)
    dialog.title("Filen finns redan")
    dialog.geometry("500x200")
    dialog.resizable(False, False)
    dialog.transient(parent)
    dialog.grab_set()

    result = None

    # Centrera på parent window
    dialog.update_idletasks()
    x = parent.winfo_x() + (parent.winfo_width() // 2) - (dialog.winfo_width() // 2)
    y = parent.winfo_y() + (parent.winfo_height() // 2) - (dialog.winfo_height() // 2)
    dialog.geometry(f"+{x}+{y}")

    # Message
    msg_frame = ttk.Frame(dialog, padding=20)
    msg_frame.pack(fill="both", expand=True)

    msg = tk.Label(
        msg_frame,
        text=f"Filen finns redan:\n\n{filepath}\n\nVad vill du göra?",
        wraplength=450,
        justify="left"
    )
    msg.pack(pady=(0, 20))

    # Buttons
    btn_frame = ttk.Frame(dialog, padding=(20, 0, 20, 20))
    btn_frame.pack(fill="x")

    def on_load():
        nonlocal result
        result = 'load'
        dialog.destroy()

    def on_update():
        nonlocal result
        result = 'update'
        dialog.destroy()

    def on_cancel():
        nonlocal result
        result = None
        dialog.destroy()

    ttk.Button(btn_frame, text="Skapa nytt protokoll", command=on_load).pack(side="left", padx=5)
    ttk.Button(btn_frame, text="Uppdatera befintligt protokoll", command=on_update).pack(side="left", padx=5)
    ttk.Button(btn_frame, text="Avbryt", command=on_cancel).pack(side="left", padx=5)

    # Hantera window close
    dialog.protocol("WM_DELETE_WINDOW", on_cancel)

    # Vänta tills dialog stängs
    parent.wait_window(dialog)

    return result


def main():
    script_dir = Path(__file__).parent.resolve()
    settings_path = script_dir / SETTINGS_FILENAME
    settings = load_settings(settings_path)

    # Normalisera output_dir tidigt (hanterar t.ex. "~" och relativa paths)
    output_dir = Path(settings.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    root = tk.Tk()
    root.title("Protokoll -> Word")

    frm = ttk.Frame(root, padding=12)
    frm.grid(row=0, column=0, sticky="nsew")
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)
    frm.columnconfigure(1, weight=1)

    # Variables (init from settings)
    team_var = tk.StringVar(value=(settings.last_team or "Teamnamn"))
    datum_var = tk.StringVar(value=str(date.today()))
    deltagare_var = tk.StringVar()
    status_var = tk.StringVar(value="På spår (🟢)")

    def add_label(row: int, text: str):
        ttk.Label(frm, text=text).grid(row=row, column=0, sticky="w", padx=(0, 10), pady=4)

    def add_entry(row: int, var: tk.StringVar):
        e = ttk.Entry(frm, textvariable=var)
        e.grid(row=row, column=1, sticky="ew", pady=4)
        return e

    add_label(0, "Team")
    add_entry(0, team_var)

    add_label(1, "Datum (YYYY-MM-DD)")
    add_entry(1, datum_var)

    add_label(2, "Deltagare (komma-separerat)")
    add_entry(2, deltagare_var)

    def add_text(row: int, title: str, height: int = 6) -> tk.Text:
        add_label(row, title)
        txt = tk.Text(frm, height=height, wrap="word")
        txt.grid(row=row, column=1, sticky="ew", pady=4)
        return txt

    vad_txt = add_text(3, "Vad vi jobbade med (en punkt per rad)", height=6)
    hinder_txt = add_text(4, "Hinder vi stötte på (en punkt per rad)", height=6)

    add_label(5, "Status")
    status_cb = ttk.Combobox(
        frm,
        textvariable=status_var,
        values=["På spår (🟢)", "Lite efter (🟡)", "Behöver hjälp (🔴)"],
        state="readonly",
    )
    status_cb.grid(row=5, column=1, sticky="w", pady=4)

    nasta_txt = add_text(6, "Nästa steg (en punkt per rad)", height=6)

    # Runtime state
    last_created_path: Path | None = None

    def validate_date_format(d: str) -> bool:
        return bool(re.match(r"^\d{4}-\d{2}-\d{2}$", (d or "").strip()))

    def compute_planned_path(team: str | None = None, datum: str | None = None) -> Path:
        t = (team if team is not None else team_var.get()).strip() or "Teamnamn"
        d = (datum if datum is not None else datum_var.get()).strip()
        return planned_docx_path(t, d, output_dir=output_dir)

    def load_existing_into_form(path: Path):
        """Ladda in innehållet från ett befintligt Word-dokument i formuläret."""
        try:
            data = load_existing_docx(path)

            # Fyll i formuläret med befintliga data
            team_var.set(data["team"])
            datum_var.set(data["datum"])
            deltagare_var.set(data["deltagare"])
            status_var.set(data["status"])

            # Rensa och fyll i text-fälten
            vad_txt.delete("1.0", "end")
            vad_txt.insert("1.0", data["vad_vi_jobbat_med"])

            hinder_txt.delete("1.0", "end")
            hinder_txt.insert("1.0", data["hinder"])

            nasta_txt.delete("1.0", "end")
            nasta_txt.insert("1.0", data["nasta_steg"])

            messagebox.showinfo(
                "Protokoll inladdat",
                f"Befintligt protokoll har laddats in i formuläret.\n\n{path}\n\n"
                "Du kan nu fortsätta redigera och spara det.",
                parent=root,
            )
        except Exception as e:
            messagebox.showerror(
                "Kunde inte ladda protokollet",
                f"Ett fel uppstod när befintligt protokoll skulle laddas:\n{e}",
                parent=root,
            )

    def check_existing_on_start():
        # Vid uppstart: använd settings.last_team (redan i team_var) + dagens datum (datum_var)
        if not validate_date_format(datum_var.get()):
            return

        path = compute_planned_path()
        if path.exists():
            load_into_form = messagebox.askyesno(
                "Befintligt protokoll hittat",
                f"Det finns redan ett protokoll med samma namn:\n\n{path}\n\n"
                "Vill du ladda in det i formuläret för att fortsätta redigera?",
                parent=root,
            )
            if load_into_form:
                load_existing_into_form(path)

    def persist_settings():
        # Spara alltid ev. uppdaterat team + ev. output_dir
        settings.output_dir = str(output_dir)
        settings.last_team = team_var.get().strip() or "Teamnamn"
        try:
            save_settings(settings_path, settings)
        except Exception:
            pass

    def on_generate():
        nonlocal last_created_path

        team = team_var.get().strip() or "Teamnamn"
        datum = datum_var.get().strip()
        deltagare = deltagare_var.get().strip()

        if not validate_date_format(datum):
            messagebox.showerror(
                "Fel",
                "Datum måste vara i formatet YYYY-MM-DD (t.ex. 2026-01-20).",
                parent=root,
            )
            return

        outpath = planned_docx_path(team, datum, output_dir=output_dir)

        if outpath.exists():
            user_choice = ask_file_exists_action(root, outpath)
            if user_choice == 'load':
                # Ladda in i formuläret
                load_existing_into_form(outpath)
                return
            elif user_choice is None:
                # Avbryt
                return
            # user_choice == 'update': fortsätt och uppdatera filen

        try:
            created = create_word_doc_to_path(
                outpath=outpath,
                team=team,
                datum=datum,
                deltagare=deltagare or "-",
                vad_vi_jobbat_med=vad_txt.get("1.0", "end"),
                hinder=hinder_txt.get("1.0", "end"),
                status_label=status_var.get(),
                nasta_steg=nasta_txt.get("1.0", "end"),
            )
        except Exception as e:
            messagebox.showerror("Fel", f"Kunde inte skapa dokumentet:\n{e}", parent=root)
            return

        last_created_path = created
        persist_settings()

        messagebox.showinfo("Klart", f"Word-dokument sparat:\n{created}", parent=root)
        try:
            open_file_with_default_app(created)
        except Exception:
            pass

    def on_upload_github():
        nonlocal last_created_path

        # Säkerställ att det finns en fil att ladda upp:
        if last_created_path is None or not last_created_path.exists():
            if validate_date_format(datum_var.get()):
                candidate = compute_planned_path()
                if candidate.exists():
                    last_created_path = candidate

        if last_created_path is None or not last_created_path.exists():
            messagebox.showwarning(
                "Ingen fil att ladda upp",
                "Skapa protokollet först (eller se till att en planerad fil finns i output-mappen).",
                parent=root,
            )
            return

        try:
            path_in_repo, branch = upload_docx_flow(root, last_created_path, settings.github)
        except Exception as e:
            messagebox.showerror("GitHub misslyckades", str(e), parent=root)
            return

        messagebox.showinfo(
            "Uppladdat",
            f"Filen är kopierad, committad och pushad:\n\n{path_in_repo}\n\nBranch: {branch}",
            parent=root,
        )

    def on_close():
        persist_settings()
        root.destroy()

    # Buttons
    btns = ttk.Frame(frm)
    btns.grid(row=7, column=0, columnspan=2, sticky="e", pady=(10, 0))

    ttk.Button(btns, text="Skapa Word-dokument", command=on_generate).grid(row=0, column=0, padx=(0, 8))
  #  ttk.Button(btns, text="Skicka till GitHub", command=on_upload_github).grid(row=0, column=1)

    root.protocol("WM_DELETE_WINDOW", on_close)

    # Kör uppstartskontroll efter att GUI laddats
    root.after(200, check_existing_on_start)

    root.mainloop()


if __name__ == "__main__":
    main()

